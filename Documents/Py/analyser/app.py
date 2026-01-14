import streamlit as st
import pandas as pd
import zipfile
import io
import json
import base64
from openai import OpenAI
from docx import Document

# ---------------------------------------------------------
# Grundopsætning (SKAL ligge øverst)
# ---------------------------------------------------------
st.set_page_config(page_title="Analyser", layout="wide")

def get_secret(key: str, default: str | None = None) -> str | None:
    if key in st.secrets:
        return str(st.secrets.get(key))
    return os.getenv(key, default)

ACCESS_KEY = get_secret("ACCESS_KEY")

if not ACCESS_KEY:
    st.error("Manglende ACCESS_KEY i Streamlit secrets.")
    st.stop()
# --- SIKKERHEDSTJEK ---
# Vi tjekker om URL'en indeholder vores hemmelige nøgle
query_params = st.query_params  # Henter parametre fra URL'en

# Hvis nøglen mangler eller er forkert, stop appen
# Vi bruger .get() så den ikke crasher hvis parameteren slet ikke findes
if query_params.get("access") != ACCESS_KEY:
    st.error("⛔ Adgang nægtet.")
    st.info("Denne app kan kun tilgås gennem Generaxions interne systemer.")
    st.stop() # Stopper koden her, så resten ikke vises

# ---------------------------------------------------------
# Selve Appen starter her
# ---------------------------------------------------------
st.title("Analyser")

# ---------------------------------------------------------
# Sidebar – modelvalg (SEO er fastlåst)
# ---------------------------------------------------------
st.sidebar.header("Indstillinger")

# Afdeling er låst til SEO i denne version
department = "SEO (Organisk)"

model_choice = st.sidebar.selectbox(
    "Vælg AI-model",
    ["Grundig (GPT-5.1)", "Hurtig (GPT-4.1)"],
    index=0
)

if "GPT-5.1" in model_choice:
    selected_model = "gpt-5.1"
else:
    selected_model = "gpt-4.1"

# ---------------------------------------------------------
# Basisinfo om kunden
# ---------------------------------------------------------
st.subheader("1. Basisinfo")

col1, col2 = st.columns(2)
with col1:
    customer_name = st.text_input("Kundenavn", placeholder="Kundenavn")
with col2:
    customer_url = st.text_input("URL", placeholder="Website")

# ---------------------------------------------------------
# Upload – Ahrefs, Screaming Frog, GSC
# ---------------------------------------------------------
st.subheader("2. Datakilder")

st.markdown("**Ahrefs – upload alle relevante rapporter**")
ahrefs_files = st.file_uploader(
    "Upload Ahrefs-rapporter (Performance, Organic Keywords, Content Gap, Referring Domains)",
    type=["csv"],
    accept_multiple_files=True,
    help=(
        "Upload alle relevante Ahrefs-eksporter her (fx 'domain_organic_perf...', 'organic_keywords...', "
        "'content_gap...', 'referring_domains...'). Appen forsøger automatisk at fordele filerne til de rigtige sektioner."
    ),
)

st.markdown("**Screaming Frog – Crawl**")
screaming_frog_file = st.file_uploader(
    "Upload Screaming Frog-crawl (Internal All / Page Titles / Word Count)",
    type=["csv", "xlsx", "xls", "zip"],
    accept_multiple_files=False,
    help="Upload enten en samlet CSV/Excel eller en ZIP med eksportfiler."
)

st.markdown("**Google Search Console – Keyword-data (valgfrit)**")
gsc_files = st.file_uploader(
    "Upload GSC Search Analytics eksport (CSV/Excel)",
    type=["csv", "xlsx", "xls"],
    accept_multiple_files=True,
    help="Fx eksport fra 'Performance' – queries/URL/clicks/impressions/position. Du kan uploade flere filer. Excel-filer med flere faner læses som én samlet datapakke, hvor hver fane gemmes separat."
)

# ---------------------------------------------------------
# Vælg slides og ekstra sektioner
# ---------------------------------------------------------
st.subheader("3. Output – Slides og ekstra sektioner")

# Bemærk: Dette er KUN til at signalere fokus til modellen.
# Den faktiske slide-struktur er låst i prompten.
slide_options = [
    "Trafik fra websitets organiske søgeord",
    "Søgeord der genererer trafik",
    "Fokus på trafikskabende organiske søgeord",
    "Organiske søgeord med uforløst potentiale",
    "Hvor vinder jeres konkurrenter?",
    "Pagetitles",
    "Antal refererende domæner til websitet",
    "EEAT",
    "Teknisk sundhedstjek (teknisk SEO)",
    "Bedre indhold",
    "Fokus",
]

selected_slides = st.multiselect(
    "Vælg hvilke temaer der skal have ekstra fokus i anbefalingerne",
    slide_options,
    default=slide_options,
)

extra_slides_text = st.text_area(
    "Ekstra slides / noter (valgfrit)",
    placeholder="Skriv korte stikord eller bullets til ekstra slides, du vil have med."
)

# Per-slide kommentarer fra rådgiveren
slide_notes = {}
# Per-slide billeder (valgfrit)
slide_images = {}

with st.expander("Tilføj kommentarer og billeder til de enkelte slides (valgfrit)"):
    st.markdown(
        "Skriv kort, hvad der er vigtigst at få med på hvert slide, og upload evt. et billede "
        "som du vil have med i analysen (fx SERP-screenshot, graf m.m.)."
    )
    for i, slide in enumerate(slide_options):
        # Titel for selve slidet
        st.markdown(f"**{slide}**")

        # To kolonner: kommentar (bred) + billede (smal)
        col_comment, col_image = st.columns([3, 1])

        with col_comment:
            slide_notes[slide] = st.text_area(
                "Kommentar",
                placeholder="Fx: Fokuser på konkurrenceprægede søgeord og fundament for hurtige resultater.",
                key=f"note_{slide}",
            )

        with col_image:
            slide_images[slide] = st.file_uploader(
                "Billede (valgfrit)",
                type=["png", "jpg", "jpeg", "webp"],
                key=f"img_{i}",
            )

        # Visuel separator mellem slides
        st.markdown("---")

# ---------------------------------------------------------
# Hjælpefunktioner til filer
# ---------------------------------------------------------
def read_tabular_file(uploaded_file):
    """Læs CSV/Excel/ZIP til en eller flere pandas DataFrames.

    Returnerer:
      - dict: {filename: df.to_dict(orient='records')}
    """
    if uploaded_file is None:
        return {}

    # ZIP med flere filer
    if uploaded_file.name.lower().endswith(".zip"):
        result = {}
        with zipfile.ZipFile(uploaded_file, "r") as z:
            for name in z.namelist():
                if name.lower().endswith(".csv"):
                    try:
                        with z.open(name) as f:
                            df = pd.read_csv(f, sep=None, engine="python")
                        result[name] = df.to_dict(orient="records")
                    except Exception as e:
                        result[name] = {"error": str(e)}
                elif name.lower().endswith((".xlsx", ".xls")):
                    try:
                        with z.open(name) as f:
                            excel_bytes = f.read()
                        excel_buf = io.BytesIO(excel_bytes)
                        xls = pd.ExcelFile(excel_buf)
                        for sheet_name in xls.sheet_names:
                            df = xls.parse(sheet_name)
                            key = f"{name}::{sheet_name}"
                            result[key] = df.to_dict(orient="records")
                    except Exception as e:
                        result[name] = {"error": str(e)}
        return result

    # Almindelig CSV/Excel
    try:
        if uploaded_file.name.lower().endswith(".csv"):
            df = pd.read_csv(uploaded_file, sep=None, engine="python")
            return {uploaded_file.name: df.to_dict(orient="records")}
        else:
            # Læs alle faner fra Excel som separate datasæt
            excel_bytes = uploaded_file.read()
            excel_buf = io.BytesIO(excel_bytes)
            xls = pd.ExcelFile(excel_buf)
            result = {}
            for sheet_name in xls.sheet_names:
                df = xls.parse(sheet_name)
                key = f"{uploaded_file.name}::{sheet_name}"
                result[key] = df.to_dict(orient="records")
            return result
    except Exception as e:
        return {uploaded_file.name: {"error": str(e)}}

def build_data_payload():
    """Samler alle uploadede filer i én struktureret data-payload."""
    data = {}

    # Ahrefs-filer (Performance, Keywords, Content Gap, Referring Domains)
    if ahrefs_files:
        data["ahrefs_performance"] = {}
        data["ahrefs_keywords_customer"] = {}
        data["ahrefs_content_gap"] = {}
        data["ahrefs_ref_domains"] = {}
        data["ahrefs_other"] = {}

        for f in ahrefs_files:
            name = f.name.lower()
            file_dict = read_tabular_file(f)

            # Performance-rapporter (fx domain_organic_perf...)
            if "perf" in name or "performance" in name:
                data["ahrefs_performance"].update(file_dict)

            # Content Gap-rapporter
            elif "content_gap" in name or "gap" in name:
                data["ahrefs_content_gap"].update(file_dict)

            # Referring domains / backlinks
            elif "referring" in name or "backlink" in name or "ref_domains" in name:
                data["ahrefs_ref_domains"].update(file_dict)

            # Organic keywords for kunden
            elif "keyword" in name or "organic" in name:
                data["ahrefs_keywords_customer"].update(file_dict)

            # Hvis vi ikke kan gætte typen, gemmes filen som "other"
            else:
                data["ahrefs_other"].update(file_dict)

    # Screaming Frog
    if screaming_frog_file:
        data["screaming_frog"] = read_tabular_file(screaming_frog_file)

    # Google Search Console (kan være flere filer)
    if gsc_files:
        data["gsc"] = {}
        for f in gsc_files:
            data["gsc"].update(read_tabular_file(f))

    return data


# ---------------------------------------------------------
# DOCX-helper: Byg DOCX fra markdown-lignende AI-output
# ---------------------------------------------------------
def build_docx_from_markdown(ai_output: str, customer_name: str = None, customer_url: str = None) -> io.BytesIO:
    """
    Bygger en DOCX-rapport ud fra det markdown-lignende output (### ..., **Anbefalinger**, bullets)
    så det visuelt matcher online-versionen bedst muligt.
    """
    doc = Document()

    # Titel
    title_text = "SEO-analyse"
    if customer_name:
        title_text += f" – {customer_name}"
    doc.add_heading(title_text, level=0)

    if customer_url:
        p_url = doc.add_paragraph()
        p_url.add_run(customer_url).italic = True

    lines = ai_output.splitlines()
    for raw_line in lines:
        line = raw_line.rstrip()
        if not line.strip():
            # spring tomme linjer over for at undgå for meget luft
            continue

        # Slide-overskrift: "### ..."
        if line.startswith("### "):
            heading_text = line[4:].strip()
            doc.add_heading(heading_text, level=1)
            continue

        # Generelle fed-overskrifter i markdown-stil, fx "**Analyse – Spor 1 (her og nu)**"
        stripped = line.strip()
        if stripped.startswith("**") and stripped.endswith("**") and len(stripped) > 4:
            bold_text = stripped[2:-2].strip()
            p = doc.add_paragraph()
            run = p.add_run(bold_text)
            run.bold = True
            continue

        # Bullets: linjer der starter med "- "
        if line.lstrip().startswith("- "):
            bullet_text = line.lstrip()[2:].strip()
            doc.add_paragraph(bullet_text, style="List Bullet")
            continue

        # Fald tilbage: almindeligt afsnit
        doc.add_paragraph(line.strip())

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ---------------------------------------------------------
# OpenAI-klient
# ---------------------------------------------------------
client = OpenAI(api_key=st.secrets["OPENAI_API_KEY"])

# ---------------------------------------------------------
# Prompt-builder (fælles for sync + streaming)
# ---------------------------------------------------------
def build_prompt(
    customer_name: str,
    customer_url: str,
    selected_slides: list,
    extra_slides_text: str,
    slide_notes_text: str,
    serialized_data: str,
) -> str:
    prompt = f"""
Du er en senior SEO-specialist og skal udarbejde en struktureret kundeanalyse,
der senere skal lægges direkte ind som tekst til slides.

Kontekst om kunden:
- Kundenavn: {customer_name or 'Ikke angivet'}
- URL: {customer_url or 'Ikke angivet'}

Rådgiveren har markeret følgende temaer som særligt vigtige at få tydelige anbefalinger på:
{json.dumps(selected_slides, ensure_ascii=False)}

Ekstra ønsker/noter fra rådgiveren:
{extra_slides_text or 'Ingen'}

Rådgiverens kommentarer til de enkelte slides (brug dem aktivt til at vinkle og prioritere indholdet på hvert slide):
{slide_notes_text or 'Ingen specifikke kommentarer til enkelte slides'}

Du modtager data i JSON-format fra:
- Ahrefs Performance (trafik, brand/non-brand, intent, osv.)
- Ahrefs Organic Keywords (kunde + evt. konkurrenter)
- Ahrefs Content Gap (konkurrent-sammenligning og manglende sider/temaer)
- Ahrefs Referring Domains / Backlinks (antal og udvikling i refererende domæner)
- Screaming Frog-crawl (titles, word count, teknisk)
- Google Search Console eksport (queries, clicks, impressions, position) hvis det findes – men analysen skal altid kunne stå alene på Ahrefs- og crawl-data.

Her er et nedklippet uddrag af data-payloaden i JSON-format (maks ca. 20.000 tegn). Du SKAL bruge dette aktivt i analysen og referere til konkrete tal, hvor det er relevant:

{serialized_data}

OPGAVE:
1) For hver slide-overskrift ovenfor skal du skrive en sektion med følgende rammer:
   - Start med "### [overskrift]" (som angivet ovenfor), så det bliver en tydelig, større overskrift i Markdown – altså uden "Slide X:" foran.
   - Skriv derefter 2–3 meget korte analyseafsnit, ikke kun anbefalinger:
     * Afsnit 1 beskriver kort, hvad data viser (konkrete tal, mønstre, udvikling, fordeling).
     * Afsnit 2 (og evt. 3) beskriver de vigtigste problemer/fejl/mangler og det største potentiale – vær meget konkret og ærlig.
   - Hvert afsnit må kun være 1–2 sætninger, og der skal være en tom linje mellem afsnittene, så teksten bliver let at kopiere direkte ind på et slide.
   - Den samlede tekst på hver slide (inkl. mellemrum) må som tommelfingerregel ikke overstige ca. 450–500 tegn. Det er vigtigere at være skarp og selektiv end udtømmende.
   - På slides 1–10 skriver du KUN analyse (ingen sektion "Anbefalinger" på disse slides).
   - På Slide 11 (Fokus) skriver du efter analysen en tydelig sektion med samlede anbefalinger for hele analysen: start med "**Anbefalinger**" (i fed) på en ny linje og skriv derefter 3–6 punktopstillede anbefalinger, der opsummerer de vigtigste næste skridt på tværs af alle slides.

2) Brug faktiske tal og mønstre fra dataen, når det er muligt. Hvis et tal ikke kan aflæses direkte, så brug kvalitative formuleringer som "lav", "mellem", "høj" fremfor at gætte procenter eller eksakte værdier. Du må ikke opfinde konkurrentnavne eller tal – brug kun navne/tal der reelt findes i dataen. Hvis data er begrænsede, skal du stadig skrive en sammenhængende analyse på hver slide baseret på de mønstre, du kan ane kombineret med generel SEO-viden – men du må ALDRIG nævne manglende data, manglende filer, værktøjer eller formuleringer som "ingen data", "materialet viser ikke", "crawlen er ikke vedlagt" eller lignende.
   På hver slide skal du tydeligt pege på 1–3 konkrete problemer/fejl/mangler og 1–3 centrale muligheder/potentialer, ikke kun generelle beskrivelser.

   Derudover må du ikke skrive om "kendskabsgrad", "brand awareness" eller lignende begreber. Du må gerne bruge forskelle mellem brand- og non-brand-søgninger til at forklare, hvilke typer søgninger der driver trafik, men du må ikke forsøge at forklare eller vurdere generel kendskabsgrad i markedet.

3) Dine anbefalinger må KUN handle om SEO-arbejde: indhold, struktur, intern linkbuilding, tekniske forbedringer, metadata/titler, CTR-forbedring i SERP og lignende. Du må IKKE anbefale PR, nyhedsbreve, betalt annoncering, SoMe-aktiviteter, offline-tiltag eller andre kanaler. Du må heller IKKE skrive anbefalinger om at forbedre datagrundlag, tracking eller rapporter (fx "brug GSC", "træk flere rapporter", "saml data", "tjek Ahrefs" osv.). Anbefalinger skal formuleres som konkrete forbedringer på kundens website og indhold – ikke som instrukser til specialisten om at hente mere data eller bruge specifikke værktøjer.
   Når du skriver om indhold, må du ikke bruge tomme formuleringer som "bedre indhold", "udbyg indhold" eller "mere relevant indhold" uden at forklare præcist, hvad der er galt med det nuværende (fx for korte tekster, manglende vigtige søgeord, duplikeret indhold, dårlig struktur, manglende FAQ osv.). Hver indholdsanbefaling skal knyttes til en konkret type fejl eller mangel.
   Formulér anbefalinger som kundeorienterede fokusområder (fx bullets med "Fokus 1: Optimer …", "Fokus 2: Udbyg …", "Fokus 3: Styrk …") fremfor direkte instrukser til specialisten (som "Gennemgå …", "Udvælg …", "Brug …"). Brug 1. person flertal ("vi") eller neutrale formuleringer om, hvad der skal arbejdes med, ikke kommandosprog.

4) Minimer brugen af fagbegreber. Brug kun et fagbegreb hvis det er nødvendigt, og forklar det kort i parentes første gang (fx "EEAT (Googles vurdering af troværdighed)"). 
   - EEAT må kun nævnes på Slide 8, og KUN som et supplement til konkrete observationer (f.eks. få referencer, tyndt indhold, manglende forfattersignaler). Brug det aldrig som en løs forklaring uden tydelig sammenhæng til data.
   - Undgå buzzwords og brede formuleringer som "relativt begrænsede E-E-A-T- og brand-signaler" uden konkret forankring i data.
   - Du må ALDRIG nævne værktøjer som Ahrefs, Google Search Console, Screaming Frog, Google Analytics eller lignende i teksten til kunden. Analysen skal fremstå som en ren kundevenlig SEO-analyse uden omtale af, hvordan den er lavet.

5) SPECIFIKKE KRAV TIL ENKELTE SLIDES:
   - Slide 2 (Søgeord der genererer trafik) skal fokusere på, at kunden ligger stærkt på centrale søgeord med høj konkurrence, hvor der er mange andre stærke domæner til stede. Forklar kort, hvordan de stærke placeringer giver et solidt fundament for hurtigere ekstra resultater og gør det oplagt at bygge videre med relaterede søgeord og long-tail-variationer. Undgå at gøre brand-søgninger til hovedpointen på dette slide – de må kun indgå som en mindre nuance.
   - Slide 3 (Fokus på trafikskabende organiske søgeord) skal, hvor data findes, pege på de vigtigste søgeord, der driver trafik, og adskille mellem brand/non-brand, hvis muligt. I anbefalingerne på denne slide skal du, hvor det er relevant, adskille "Spor 1 (her og nu)" for hurtige gevinster på kategorier/produktsider og "Spor 2 (langsigtet)" for guides/opskrifter og mere langsigtet indholdsopbygning.
   - Slide 4 (Organiske søgeord med uforløst potentiale) skal, hvor data findes, aktivt bruge Ahrefs Organic Keywords + Content Gap til at pege på 3–5 konkrete temaer/typer søgninger med stort potentiale (høj volume, lavere position, manglende landingssider). Nævn disse temaer eksplicit i analysen som keyword-klynger, ikke kun som generelle idéer.
   - Slide 5 (Hvor vinder jeres konkurrenter?) skal, hvor data findes, fokusere på tydelige mønstre fra Ahrefs Performance + Content Gap: hvilke emner/kategorier konkurrenter dominerer, og hvor kunden mangler indhold. Peg på 3–5 konkrete emneområder eller sider, hvor konkurrenter får betydelig trafik og kunden ikke har en tilsvarende stærk side. Brug kun navngivne brands (fx supermarkedskæder eller producentnavne), hvis det tydeligt understøtter pointen – ellers tal om "større kæder" eller "andre brands" i generelle termer.
   - Slide 6 (Pagetitles) skal altid indeholde mindst 1–2 konkrete "før/efter"-eksempler på sidetitler: én linje der starter med "Nuværende:" efterfulgt af en eksisterende titel, og én linje der starter med "Foreslået:" med en forbedret, mere sælgende titel. Det gør anbefalingerne operationelle.
   - Slide 7 (Antal refererende domæner til websitet) skal bruge faktiske tal, hvis de findes i dataen. Hvis tal ikke findes, skal du stadig skrive en generel, kundevenlig vurdering af linkstyrke og behov for flere relevante links – uden at nævne manglende data. Under anbefalinger skal du altid komme med 2–3 meget konkrete idéer til linkbuilding-tiltag (fx typer sites der kan kontaktes, konkrete indholdsidéer der kan tiltrække links), ikke kun generelle udsagn som "skab linkværdigt indhold".
   - Slide 8 (EEAT) skal være kort og konkret: 1 sætning der forklarer, hvordan EEAT ser ud lige nu, og 2–3 meget konkrete SEO-tiltag der kan styrke det (fx udfoldede kategoritekster, forfatterprofiler, case-sider, eksterne omtaler).
   - Slide 9 (Teknisk sundhedstjek (teknisk SEO)) skal, hvor data fra Screaming Frog findes, kommentere kort på tekniske forhold som tynde sider, duplikerede titler, åbenlyse 404/redirect-problemer, URL-struktur og intern linkdybde. Du må ikke digte om Core Web Vitals eller andre performance-metrics, hvis der ikke er konkret data. Hold fokus på det, der kan aflæses fra crawlen.
   - Slide 11 (Fokus) skal samle de vigtigste fokusområder og anbefalinger for de næste 3–6 måneder i et meget skarpt prioriteret format:
     * 1 kort sætning der beskriver det overordnede fokus.
     * Under sektionen "**Anbefalinger**" skal du skrive 3–6 bullets, som hver beskriver et klart fokusområde eller indsats (fx "Fokus 1: Optimer …", "Fokus 2: Udbyg …", "Fokus 3: Styrk …"). Hver bullet skal være formuleret som et kundeorienteret fokusområde, ikke en teknisk to-do. Undgå at alle bullets starter ens; variér formuleringerne, og brug primært korte beskrivelser som "Fokus X: [indsats]" fremfor at gentage "Vi anbefaler, at der arbejdes med …" i hver bullet.
5b) Analysen er 100 % kundevendt. Læseren er kunden. Du må aldrig kommentere på selve analysen, datakvaliteten eller foreslå, hvordan fremtidige analyser kan blive bedre. Ingen meta-kommentarer om processen – kun konklusioner og anbefalinger, som kunden direkte kan handle på.

6) Hold tonen professionel, direkte og uden fyldord. Du skriver til en marketingansvarlig, der forstår det grundlæggende i SEO, men ikke nødvendigvis arbejder i værktøjerne dagligt. Skriv kort, konkret og uden unødige sidespor.

7) Skriv ALTING på dansk.

8) På Slide 11 (Fokus), hvor du samler anbefalingerne, skal du så vidt muligt strukturere bullets som en lille handlingsplan: 2–3 konkrete ændringer der kan laves nu på eksisterende sider, og 1–2 forslag til nyt indhold eller tekniske indsatser, der kan bygges senere. Undgå rene floskler – hver anbefaling skal kunne omsættes direkte til en opgave i et backlog, og formuleres som et fokusområde for kunden (fx "Fokus 1: Optimer …", "Fokus 2: Udbyg …") fremfor som direkte instrukser til specialisten.

9) Anbefalinger må KUN skrives i sektionen "**Anbefalinger**" under overskriften "Fokus". På alle andre slides (1–10) må du ikke skrive sætninger der starter med "Vi anbefaler", "Fokus X:" eller på anden måde beskriver konkrete næste skridt eller indsatsområder – disse slides er udelukkende analyserende og må kun beskrive, hvad data viser, hvilke problemer der findes, og hvor potentialet ligger.

Returnér svaret som ren tekst i den viste rækkefølge, startende direkte med den første overskrift (fx "### Trafik fra websitets organiske søgeord") og uden ekstra indledning eller afsluttende kommentar.
"""
    return prompt

# ---------------------------------------------------------
# Synkront AI-kald (beholdes som fallback)
# ---------------------------------------------------------
def ask_ai(
    department: str,
    customer_name: str,
    customer_url: str,
    selected_slides: list,
    extra_slides_text: str,
    slide_notes: dict,
    slide_images: dict,
    data_payload: dict,
):
    # Vi klipper payload ned for at undgå alt for lange prompts
    serialized_data = json.dumps(data_payload, default=str)[:20000]
    slide_notes_text = ""
    if slide_notes:
        lines = []
        for slide, note in slide_notes.items():
            note_clean = (note or "").strip()
            if not note_clean:
                note_clean = "Ingen specifik kommentar."
            lines.append(f"- {slide}: {note_clean}")
        slide_notes_text = "\n".join(lines)
    prompt = build_prompt(
        customer_name=customer_name,
        customer_url=customer_url,
        selected_slides=selected_slides,
        extra_slides_text=extra_slides_text,
        slide_notes_text=slide_notes_text,
        serialized_data=serialized_data,
    )

    # Byg multimodal content med tekst + evt. billeder
    content = [
        {
            "type": "input_text",
            "text": prompt,
        }
    ]

    # Tilføj billeder pr. slide, hvis der er uploadet nogen
    if slide_images:
        for slide, uploaded_img in slide_images.items():
            if uploaded_img is None:
                continue
            try:
                img_bytes = uploaded_img.getvalue()
                b64_img = base64.b64encode(img_bytes).decode("utf-8")
                # Først lidt kontekst-tekst, så modellen ved hvilket slide billedet hører til
                content.append(
                    {
                        "type": "input_text",
                        "text": f"Billede til slide '{slide}'. Brug dette billede som ekstra kontekst i din analyse af det tilhørende tema.",
                    }
                )
                # Selve billedet (som data-URL til Responses API)
                mime = uploaded_img.type or "image/png"
                data_url = f"data:{mime};base64,{b64_img}"
                content.append(
                    {
                        "type": "input_image",
                        "image_url": data_url,
                    }
                )
            except Exception:
                # Hvis noget går galt med et enkelt billede, ignorerer vi det og fortsætter
                continue

    response = client.responses.create(
        model=selected_model,
        input=[
            {
                "role": "user",
                "content": content,
            }
        ],
    )

    try:
        return response.output_text
    except AttributeError:
        if hasattr(response, "output") and response.output and hasattr(response.output[0], "content"):
            parts = []
            for c in response.output[0].content:
                if getattr(c, "type", "") in ("output_text", "text"):
                    parts.append(getattr(c, "text", ""))
            return "\n".join(parts)
        return "Der opstod en fejl ved læsning af AI-svaret."

# ---------------------------------------------------------
# Streaming AI-kald
# ---------------------------------------------------------
def ask_ai_stream(
    department: str,
    customer_name: str,
    customer_url: str,
    selected_slides: list,
    extra_slides_text: str,
    slide_notes: dict,
    slide_images: dict,
    data_payload: dict,
):
    """Streaming-version af AI-kaldet – yield'er tekststumper løbende."""
    serialized_data = json.dumps(data_payload, default=str)[:20000]
    slide_notes_text = ""
    if slide_notes:
        lines = []
        for slide, note in slide_notes.items():
            note_clean = (note or "").strip()
            if not note_clean:
                note_clean = "Ingen specifik kommentar."
            lines.append(f"- {slide}: {note_clean}")
        slide_notes_text = "\n".join(lines)
    prompt = build_prompt(
        customer_name=customer_name,
        customer_url=customer_url,
        selected_slides=selected_slides,
        extra_slides_text=extra_slides_text,
        slide_notes_text=slide_notes_text,
        serialized_data=serialized_data,
    )

    # Byg multimodal content med tekst + evt. billeder
    content = [
        {
            "type": "input_text",
            "text": prompt,
        }
    ]

    # Tilføj billeder pr. slide, hvis der er uploadet nogen
    if slide_images:
        for slide, uploaded_img in slide_images.items():
            if uploaded_img is None:
                continue
            try:
                img_bytes = uploaded_img.getvalue()
                b64_img = base64.b64encode(img_bytes).decode("utf-8")
                content.append(
                    {
                        "type": "input_text",
                        "text": f"Billede til slide '{slide}'. Brug dette billede som ekstra kontekst i din analyse af det tilhørende tema.",
                    }
                )
                mime = uploaded_img.type or "image/png"
                data_url = f"data:{mime};base64,{b64_img}"
                content.append(
                    {
                        "type": "input_image",
                        "image_url": data_url,
                    }
                )
            except Exception:
                continue

    with client.responses.stream(
        model=selected_model,
        input=[
            {
                "role": "user",
                "content": content,
            }
        ],
    ) as stream:
        for event in stream:
            try:
                # Responses streaming events: vi går efter output_text.delta events
                if hasattr(event, "type") and event.type == "response.output_text.delta":
                    delta_text = getattr(event, "delta", None)
                    if delta_text:
                        yield str(delta_text)
            except Exception:
                # Ignorer events vi ikke kan parse – fortsæt streaming
                continue

# ---------------------------------------------------------
# Kør analyse (med streaming)
# ---------------------------------------------------------
st.subheader("4. Kør analyse")

run_analysis = st.button("Kør analyse")

ai_output = None

if run_analysis:
    # Minimal validering – vi kræver som minimum Ahrefs-data
    if not ahrefs_files:
        st.error("Du skal som minimum uploade Ahrefs-rapporter (Performance og Organic Keywords for kunden).")
    else:
        data_payload = build_data_payload()

        placeholder = st.empty()
        status = st.empty()
        status.write("Analyserer data med AI (streaming)...")

        full_text = ""

        try:
            for chunk in ask_ai_stream(
                department=department,
                customer_name=customer_name,
                customer_url=customer_url,
                selected_slides=selected_slides,
                extra_slides_text=extra_slides_text,
                slide_notes=slide_notes,
                slide_images=slide_images,
                data_payload=data_payload,
            ):
                full_text += chunk
                placeholder.markdown("### Resultat\n\n" + full_text)
        except Exception as e:
            status.empty()
            st.error(f"Der opstod en fejl i AI-streamingen: {e}")
        else:
            status.empty()
            if full_text.strip():
                st.success("Analyse gennemført.")
                ai_output = full_text
            else:
                st.error("Der opstod en fejl i AI-svaret. Prøv igen.")

# ---------------------------------------------------------
# DOCX-download
# ---------------------------------------------------------
if "ai_output" not in st.session_state:
    st.session_state["ai_output"] = None

# Gem output i session, hvis der er kommet nyt
if ai_output:
    st.session_state["ai_output"] = ai_output

if st.session_state.get("ai_output"):
    st.markdown("---")
    st.subheader("Download rapport")

    docx_buffer = build_docx_from_markdown(
        ai_output=st.session_state["ai_output"],
        customer_name=customer_name,
        customer_url=customer_url,
    )

    st.download_button(
        label="Download DOCX-rapport",
        data=docx_buffer,
        file_name=f"SEO_analyse_{(customer_name or 'kunde').replace(' ', '_')}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
